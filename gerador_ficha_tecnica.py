from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path


# ==========================================
# CONFIGURAÇÕES GERAIS
# ==========================================

ARQUIVO_SAIDA = "Ficha_Tecnica_John_Cunningham_Gerada.docx"

# Se quiser usar imagens, ajuste os caminhos abaixo.
IMG_TS0456_PRINCIPAL = "ts0456_principal.png"
IMG_TS0456_VAR2 = "ts0456_var2.png"
IMG_TS0456_VAR3 = "ts0456_var3.png"

IMG_TS0458_PRINCIPAL = "ts0458_principal.png"
IMG_TS0458_VAR2 = "ts0458_var2.png"
IMG_TS0458_VAR3 = "ts0458_var3.png"


# ==========================================
# FUNÇÕES AUXILIARES
# ==========================================

def set_page_layout(doc: Document):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, **kwargs):
    """
    Exemplo:
    set_cell_border(
        cell,
        top={"val": "single", "sz": 8, "color": "000000"},
        bottom={"val": "single", "sz": 8, "color": "000000"},
        left={"val": "single", "sz": 8, "color": "000000"},
        right={"val": "single", "sz": 8, "color": "000000"},
    )
    """
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()

    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge in ("left", "top", "right", "bottom"):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = f"w:{edge}"
            element = tc_borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tc_borders.append(element)

            element.set(qn("w:val"), edge_data.get("val", "single"))
            element.set(qn("w:sz"), str(edge_data.get("sz", 8)))
            element.set(qn("w:color"), edge_data.get("color", "000000"))
            element.set(qn("w:space"), str(edge_data.get("space", 0)))


def clear_cell(cell):
    cell.text = ""
    p = cell.paragraphs[0]
    p.clear()


def set_cell_text(
    cell,
    text: str,
    bold: bool = False,
    size: int = 9,
    color: str = "000000",
    align=WD_ALIGN_PARAGRAPH.LEFT,
    font_name: str = "Arial",
):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_run_to_paragraph(
    paragraph,
    text: str,
    bold: bool = False,
    size: int = 9,
    color: str = "000000",
    font_name: str = "Arial",
):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.color.rgb = RGBColor.from_string(color)
    return run


def add_paragraph_to_cell(
    cell,
    text: str,
    bold: bool = False,
    size: int = 9,
    color: str = "000000",
    align=WD_ALIGN_PARAGRAPH.LEFT,
    font_name: str = "Arial",
):
    p = cell.add_paragraph()
    p.alignment = align
    add_run_to_paragraph(
        p, text=text, bold=bold, size=size, color=color, font_name=font_name
    )
    return p


def set_table_borders(table):
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"val": "single", "sz": 6, "color": "000000"},
                bottom={"val": "single", "sz": 6, "color": "000000"},
                left={"val": "single", "sz": 6, "color": "000000"},
                right={"val": "single", "sz": 6, "color": "000000"},
            )


def set_col_width(cell, width_cm: float):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_w = OxmlElement("w:tcW")
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))  # aproximação para twips
    tc_w.set(qn("w:type"), "dxa")
    tc_pr.append(tc_w)


def style_black_header(cell, text: str):
    set_cell_shading(cell, "000000")
    set_cell_text(
        cell,
        text,
        bold=True,
        size=9,
        color="FFFFFF",
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )


def style_yellow_highlight(cell, text: str):
    set_cell_shading(cell, "F2C94C")
    set_cell_text(
        cell,
        text,
        bold=True,
        size=10,
        color="000000",
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )


def style_gray_cell(cell, text: str = "", bold: bool = False, align=WD_ALIGN_PARAGRAPH.LEFT):
    set_cell_shading(cell, "EDEDED")
    set_cell_text(cell, text, bold=bold, size=8, color="000000", align=align)


def add_spacer(doc, size_pt=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("")
    run.font.size = Pt(size_pt)


def try_add_image(cell, image_path: str, width_cm: float = 4.5, fallback_text: str = "[imagem]"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if Path(image_path).exists():
        run = p.add_run()
        run.add_picture(image_path, width=Cm(width_cm))
    else:
        add_run_to_paragraph(p, fallback_text, bold=False, size=8, color="666666")


def add_color_block_text(cell, title_lines, color_lines):
    cell.text = ""
    for i, line in enumerate(title_lines):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_run_to_paragraph(p, line, bold=True if i == 0 else False, size=8)

    for color_text in color_lines:
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_run_to_paragraph(p, f"■ {color_text}", size=8)


# ==========================================
# PÁGINA 1
# ==========================================

def criar_cabecalho_pagina_1(doc: Document):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)

    set_col_width(table.cell(0, 0), 7.0)
    set_col_width(table.cell(0, 1), 11.0)

    left = table.cell(0, 0)
    right = table.cell(0, 1)

    set_cell_shading(left, "000000")
    set_cell_text(left, "JOHN CUNNINGHAM", bold=True, size=14, color="FFFFFF")

    right.text = ""
    p1 = right.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_to_paragraph(
        p1,
        "FICHA DE ACOMPANHAMENTO DE PRODUÇÃO",
        bold=True,
        size=12,
    )
    p2 = right.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run_to_paragraph(p2, "AUTORIZADO POR: ____________________", size=9)


def criar_bloco_identificacao(doc: Document):
    table = doc.add_table(rows=4, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)

    larguras = [4.4, 5.2, 4.0, 4.2]
    for c in range(4):
        set_col_width(table.cell(0, c), larguras[c])

    # Linha 1
    set_cell_text(table.cell(0, 0), "REFERÊNCIA: TS0454 A TS0458", bold=True, size=8)
    set_cell_text(table.cell(0, 1), "PRODUTO: T-SHIRTS ESTAMPADAS", bold=True, size=8)
    set_cell_text(table.cell(0, 2), "[ ] FICHA 1 (ACOMPANHAMENTO)", size=8)
    set_cell_text(table.cell(0, 3), "[ ] FICHA 2 (PRODUÇÃO)", size=8)

    # Linha 2
    set_cell_text(table.cell(1, 0), "OP.: ____________________", size=8)
    set_cell_text(table.cell(1, 1), "GRADE: 1 - 2 - 2 - 1", size=8)
    set_cell_text(table.cell(1, 2), "P   |   M   |   G   |   GG   |   TOTAL", size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table.cell(1, 3), "LANÇADA POR: ____________________", size=8)

    # Linha 3
    set_cell_text(table.cell(2, 0), "DESCRIÇÃO: MENEGOTTI 100% ALGODÃO (MAXI)", size=8)
    set_cell_text(table.cell(2, 1), "DESIGNER: MECO VENTURA", size=8)
    set_cell_text(table.cell(2, 2), "DATA: ____ / ____ / ______", size=8)
    set_cell_text(table.cell(2, 3), "", size=8)

    # Linha 4
    merged = table.cell(3, 0).merge(table.cell(3, 3))
    style_yellow_highlight(merged, "COLEÇÃO: OUTONO DROP.1")


def criar_bloco_costura(container_cell):
    container_cell.text = ""
    outer = container_cell.add_table(rows=8, cols=1)
    outer.alignment = WD_TABLE_ALIGNMENT.CENTER
    outer.autofit = False
    set_table_borders(outer)

    style_black_header(outer.cell(0, 0), "COSTURA")

    style_gray_cell(outer.cell(1, 0), "DATA DE ENTREGA: ____ / ____ / ______    |    PREV. RETORNO: ____ / ____ / ______", align=WD_ALIGN_PARAGRAPH.LEFT)

    style_gray_cell(outer.cell(2, 0), "P   |   M   |   G   |   GG   |   TOTAL", align=WD_ALIGN_PARAGRAPH.CENTER)

    style_gray_cell(outer.cell(3, 0), "ENTREGUE: ____________________    |    RECEBIDO: ____________________", align=WD_ALIGN_PARAGRAPH.LEFT)

    style_gray_cell(outer.cell(4, 0), "CONFERÊNCIA DO MATERIAL NA DATA: ____ / ____ / ______", align=WD_ALIGN_PARAGRAPH.LEFT)

    # Tabela de defeitos dentro da linha 5
    defeitos_cell = outer.cell(5, 0)
    defeitos_cell.text = ""
    defeitos = defeitos_cell.add_table(rows=6, cols=5)
    defeitos.alignment = WD_TABLE_ALIGNMENT.CENTER
    defeitos.autofit = False
    set_table_borders(defeitos)

    headers = ["DEFEITOS", "P", "M", "G", "GG"]
    for i, h in enumerate(headers):
        style_gray_cell(defeitos.cell(0, i), h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    linhas = ["TECIDO", "OUTROS", "COSTURA", "ESTAMPARIA", "SEM PRODUZIR"]
    for r, nome in enumerate(linhas, start=1):
        style_gray_cell(defeitos.cell(r, 0), nome, align=WD_ALIGN_PARAGRAPH.LEFT)
        for c in range(1, 5):
            style_gray_cell(defeitos.cell(r, c), "", align=WD_ALIGN_PARAGRAPH.CENTER)

    style_gray_cell(outer.cell(6, 0), "EMISSÃO DE RECIBO - DATA: ____ / ____ / ______", align=WD_ALIGN_PARAGRAPH.LEFT)
    style_gray_cell(outer.cell(7, 0), "OBS.:", align=WD_ALIGN_PARAGRAPH.LEFT)


def criar_bloco_checklist(container_cell):
    container_cell.text = ""
    outer = container_cell.add_table(rows=2, cols=1)
    outer.alignment = WD_TABLE_ALIGNMENT.CENTER
    outer.autofit = False
    set_table_borders(outer)

    style_black_header(outer.cell(0, 0), "CHECK LIST")

    body = outer.cell(1, 0)
    body.text = ""
    table = body.add_table(rows=16, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)

    headers = ["ITEM", "QTD.", "SIM", "NÃO", "ARTIGO"]
    for i, h in enumerate(headers):
        style_gray_cell(table.cell(0, i), h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    itens = [
        "MOLDES",
        "FRENTE",
        "COSTAS",
        "MANGAS",
        "AVIAMENTOS",
        "CÓD. BARRAS",
        "LACRE",
        "TAG",
        "SACOLA",
        "ETIQ. COMPOSIÇÃO",
        "ETIQ. TAMANHO",
        "ETIQUETA EXTERNA",
        "ETIQUETA INTERNA",
        "TERMOCOLANTE",
        "RIBANA",
    ]

    for r, item in enumerate(itens, start=1):
        style_gray_cell(table.cell(r, 0), item, align=WD_ALIGN_PARAGRAPH.LEFT)
        for c in range(1, 5):
            style_gray_cell(table.cell(r, c), "", align=WD_ALIGN_PARAGRAPH.CENTER)

    # adiciona rodapé abaixo
    add_paragraph_to_cell(body, "SEPARADO POR: ____________________", size=8)
    add_paragraph_to_cell(body, "CONFERIDO POR: ____________________", size=8)
    add_paragraph_to_cell(body, "VIÉS", size=8)


def criar_bloco_estoque(container_cell):
    container_cell.text = ""
    outer = container_cell.add_table(rows=4, cols=1)
    outer.alignment = WD_TABLE_ALIGNMENT.CENTER
    outer.autofit = False
    set_table_borders(outer)

    style_black_header(outer.cell(0, 0), "ESTOQUE")
    style_gray_cell(outer.cell(1, 0), "ENTRADA DE PEÇAS", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    style_gray_cell(outer.cell(2, 0), "P   |   M   |   G   |   GG   |   TOTAL", align=WD_ALIGN_PARAGRAPH.CENTER)
    style_gray_cell(outer.cell(3, 0), "CONFERIDO POR: ____________________    |    OBS.:", align=WD_ALIGN_PARAGRAPH.LEFT)


def criar_corpo_pagina_1(doc: Document):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    set_col_width(table.cell(0, 0), 11.5)
    set_col_width(table.cell(0, 1), 7.5)

    left = table.cell(0, 0)
    right = table.cell(0, 1)

    criar_bloco_costura(left)

    right.text = ""
    wrapper = right.add_table(rows=2, cols=1)
    wrapper.alignment = WD_TABLE_ALIGNMENT.CENTER
    wrapper.autofit = False
    set_table_borders(wrapper)

    criar_bloco_checklist(wrapper.cell(0, 0))
    criar_bloco_estoque(wrapper.cell(1, 0))


# ==========================================
# PÁGINA 2
# ==========================================

def criar_tabela_estamparia(container_cell, linhas):
    container_cell.text = ""
    table = container_cell.add_table(rows=len(linhas) + 1, cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)

    headers = ["CORPO", "TELA - 1", "TELA - 2", "TELA - 3", "TELA - 4", "TELA - 5", "TELA - ETIQUETA"]
    for i, h in enumerate(headers):
        style_black_header(table.cell(0, i), h)

    for r, linha in enumerate(linhas, start=1):
        for c, valor in enumerate(linha):
            style_gray_cell(table.cell(r, c), valor, align=WD_ALIGN_PARAGRAPH.CENTER)


def criar_bloco_visual_produto(
    doc: Document,
    referencia: str,
    produto: str,
    principal_img: str,
    var2_img: str,
    var3_img: str,
    bloco_esquerdo,
    bloco_direito_1,
    bloco_direito_2,
    tabela_linhas,
):
    # Cabeçalho do bloco
    head = doc.add_table(rows=1, cols=2)
    head.alignment = WD_TABLE_ALIGNMENT.CENTER
    head.autofit = False
    set_table_borders(head)

    set_col_width(head.cell(0, 0), 10.5)
    set_col_width(head.cell(0, 1), 8.0)

    set_cell_text(head.cell(0, 0), f"REFERÊNCIA: {referencia}", bold=True, size=9)
    set_cell_text(head.cell(0, 1), f"PRODUTO: {produto}", bold=True, size=9)

    # Área principal
    body = doc.add_table(rows=1, cols=3)
    body.alignment = WD_TABLE_ALIGNMENT.CENTER
    body.autofit = False
    set_table_borders(body)

    set_col_width(body.cell(0, 0), 5.0)
    set_col_width(body.cell(0, 1), 7.5)
    set_col_width(body.cell(0, 2), 6.0)

    # Esquerda
    add_color_block_text(body.cell(0, 0), bloco_esquerdo["titulos"], bloco_esquerdo["cores"])

    # Centro
    centro = body.cell(0, 1)
    centro.text = ""
    try_add_image(centro, principal_img, width_cm=5.8, fallback_text="[imagem principal]")
    add_paragraph_to_cell(centro, "SILK SCREEN ESTAMPA INTERNA CENTRALIZADA NAS COSTAS", size=8, color="C00000", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph_to_cell(centro, "GOLA EM RIBANA (4,3 CM ABERTA)", size=8, color="C00000", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph_to_cell(centro, "ABANADO (2,1 CM)", size=8, color="C00000", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Direita
    direita = body.cell(0, 2)
    direita.text = ""
    subt = direita.add_table(rows=2, cols=2)
    subt.alignment = WD_TABLE_ALIGNMENT.CENTER
    subt.autofit = False
    set_table_borders(subt)

    add_color_block_text(subt.cell(0, 0), bloco_direito_1["titulos"], bloco_direito_1["cores"])
    try_add_image(subt.cell(0, 1), var2_img, width_cm=3.0, fallback_text="[var.02]")

    add_color_block_text(subt.cell(1, 0), bloco_direito_2["titulos"], bloco_direito_2["cores"])
    try_add_image(subt.cell(1, 1), var3_img, width_cm=3.0, fallback_text="[var.03]")

    # Tabela estamparia
    criar_tabela_estamparia(doc.add_table(rows=1, cols=1).cell(0, 0), tabela_linhas)

    p_obs = doc.add_paragraph()
    p_obs.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run_to_paragraph(p_obs, "OBS.: TINTA COMUM.", bold=True, size=8, color="C00000")
    p_obs2 = doc.add_paragraph()
    p_obs2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run_to_paragraph(p_obs2, "FAZER AS CORES O MAIS APROXIMADO POSSÍVEL.", bold=True, size=8, color="C00000")


# ==========================================
# GERAÇÃO DO DOCUMENTO
# ==========================================

def gerar_documento():
    doc = Document()
    set_page_layout(doc)

    # Página 1
    criar_cabecalho_pagina_1(doc)
    add_spacer(doc, 3)
    criar_bloco_identificacao(doc)
    add_spacer(doc, 4)
    criar_corpo_pagina_1(doc)

    # Quebra de página
    doc.add_page_break()

    # Bloco superior - TS0456
    criar_bloco_visual_produto(
        doc=doc,
        referencia="TS0456",
        produto="BUILT DIFFERENT",
        principal_img=IMG_TS0456_PRINCIPAL,
        var2_img=IMG_TS0456_VAR2,
        var3_img=IMG_TS0456_VAR3,
        bloco_esquerdo={
            "titulos": [
                "(VAR.01) - TALCO (8006)",
                "SILK CINZA BASE (INTERNO)",
                "LINHA NA COR DO TECIDO",
            ],
            "cores": [
                "ESTAMPA EXTERNA (FRENTE) COR 1 - CINZA CLARO",
                "ESTAMPA EXTERNA (FRENTE) COR 2 - BEGE",
                "ESTAMPA EXTERNA (FRENTE) COR 3 - AZUL (MALHA)",
                "ESTAMPA EXTERNA (FRENTE) COR 4 - PRETO CROMIA",
            ],
        },
        bloco_direito_1={
            "titulos": [
                "(VAR.02) - STAR DUST (2101)",
                "SILK BRANCO (INTERNO)",
                "LINHA NA COR DO TECIDO",
            ],
            "cores": [
                "COR 1 - TOM SOBRE TOM",
                "COR 2 - BEGE",
                "COR 3 - OFF-WHITE",
                "COR 4 - OFF-WHITE",
            ],
        },
        bloco_direito_2={
            "titulos": [
                "(VAR.03) - PRETO (6500)",
                "SILK BRANCO (INTERNO)",
                "LINHA NA COR DO TECIDO",
            ],
            "cores": [
                "COR 1 - CHUMBO",
                "COR 2 - BEGE",
                "COR 3 - AZUL (MALHA)",
                "COR 4 - OFF-WHITE",
            ],
        },
        tabela_linhas=[
            ["TALCO (8006)", "CINZA CLARO", "BEGE", "AZUL (MALHA)", "PRETO CROMIA", "", "CINZA BASE"],
            ["STAR DUST (2101)", "TOM SOBRE TOM", "BEGE", "OFF-WHITE", "OFF-WHITE", "", "BRANCO 80%"],
            ["PRETO (6500)", "CHUMBO", "BEGE", "AZUL (MALHA)", "OFF-WHITE", "", "BRANCO 80%"],
        ],
    )

    add_spacer(doc, 10)

    # Bloco inferior - TS0458
    criar_bloco_visual_produto(
        doc=doc,
        referencia="TS0458",
        produto="LOST IN",
        principal_img=IMG_TS0458_PRINCIPAL,
        var2_img=IMG_TS0458_VAR2,
        var3_img=IMG_TS0458_VAR3,
        bloco_esquerdo={
            "titulos": [
                "(VAR.01) - TALCO (8006)",
                "SILK CINZA BASE (INTERNO)",
                "LINHA NA COR DO TECIDO",
            ],
            "cores": [
                "ESTAMPA EXTERNA (FRENTE) COR 1 - NÃO USAR",
                "COR 2 - PRETO CROMIA",
                "COR 3 - VERMELHO (PLASTISOL)",
            ],
        },
        bloco_direito_1={
            "titulos": [
                "(VAR.02) - CASTANHA AM (2406)",
                "SILK BRANCO (INTERNO)",
                "LINHA NA COR DO TECIDO",
            ],
            "cores": [
                "COR 1 - OFF-WHITE",
                "COR 2 - PRETO CROMIA",
                "COR 3 - VERMELHO (PLASTISOL)",
            ],
        },
        bloco_direito_2={
            "titulos": [
                "(VAR.03) - PRETO (6500)",
                "SILK BRANCO (INTERNO)",
                "LINHA NA COR DO TECIDO",
            ],
            "cores": [
                "COR 1 - OFF-WHITE",
                "COR 2 - PRETO CROMIA",
                "COR 3 - VERMELHO (PLASTISOL)",
            ],
        },
        tabela_linhas=[
            ["TALCO (8006)", "NÃO USAR", "PRETO CROMIA", "VERMELHO (PLAST.)", "", "", "CINZA BASE"],
            ["CASTANHA AM (2406)", "OFF-WHITE", "PRETO CROMIA", "VERMELHO (PLAST.)", "", "", "BRANCO 80%"],
            ["PRETO (6500)", "OFF-WHITE", "PRETO CROMIA", "VERMELHO (PLAST.)", "", "", "BRANCO 80%"],
        ],
    )

    doc.save(ARQUIVO_SAIDA)
    print(f"Documento gerado com sucesso: {ARQUIVO_SAIDA}")


if __name__ == "__main__":
    gerar_documento()